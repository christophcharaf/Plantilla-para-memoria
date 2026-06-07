#!/usr/bin/env python3
"""Completa el Informe de avance.docx con datos del proyecto real."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from task_grid_table import replace_task_table_after_anchor

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Avance" / "Informe de avance.docx"

GREEN = "00FF00"
YELLOW = "FFFF00"
RED = "FF0000"

# (nombre, color, costo, tiempo, critico) — solo subtareas numeradas (formato grilla)
TASKS = [
    ("1.1 Revisión del estado del arte", GREEN, "$", "=", False),
    ("1.2 Análisis de métricas en Prometheus", GREEN, "$", "=", True),
    ("1.3 Definición de requerimientos", GREEN, "$", "=", False),
    ("1.4 Diseño preliminar del pipeline", GREEN, "$", "=", True),
    ("2.1 Recolector de métricas desde Prometheus", GREEN, "$", "=", True),
    ("2.2 Preprocesamiento y normalización", GREEN, "$", "=", True),
    ("2.3 Ingeniería de variables temporales", GREEN, "$", "=", False),
    ("2.4 Ventanas deslizantes", GREEN, "$", "=", True),
    ("2.5 Validación del pipeline", GREEN, "$", "=", False),
    ("3.1 Definición arquitectura LSTM Autoencoder", GREEN, "$", "=", True),
    ("3.2 Implementación en Keras/TensorFlow", GREEN, "$", "=", False),
    ("3.3 Preparación conjuntos train/val", GREEN, "$", "=", False),
    ("3.4 Entrenamiento y tuning", GREEN, "$", "+", True),
    ("3.5 Evaluación con métricas objetivas", GREEN, "$", "=", False),
    ("4.1 Mecanismo de alerta con umbral", GREEN, "$", "=", True),
    ("4.2 Integración con Opsgenie", GREEN, "$", "=", True),
    ("4.3 Enlaces a dashboards Grafana", GREEN, "$", "=", False),
    ("4.4 Pruebas de integración", GREEN, "$", "=", False),
    ("5.1 Validación funcional con operaciones", GREEN, "$", "=", True),
    ("5.2 Informe de resultados y validación", GREEN, "$", "=", False),
    ("5.3 Redacción memoria (1° bimestre)", GREEN, "$", "=", False),
    ("5.4 Redacción memoria (2° bimestre)", YELLOW, "$", "+", True),
    ("5.5 Documentación de uso e integración", GREEN, "$", "=", False),
    ("5.6 Informe de avances", GREEN, "$", "=", False),
    ("5.7 Preparación presentación final", YELLOW, "$", "=", False),
]

REQUIREMENTS = [
    ("Req 1.1: Recolectar métricas técnicas desde Prometheus vía API REST.", GREEN),
    ("Req 1.2: Construir ventanas deslizantes de 20 pasos (10 min) con muestreo de 30 s.", GREEN),
    ("Req 1.3: Aplicar normalización y codificación cíclica temporal (hour_sin, hour_cos, etc.).", GREEN),
    ("Req 1.4: Reconstruir secuencias multivariadas con el autoencoder LSTM.", GREEN),
    ("Req 1.5: Detectar anomalías cuando el error de reconstrucción supere un umbral configurable.", GREEN),
    ("Req 1.6: Enviar alertas automáticas a través de Opsgenie.", GREEN),
    ("Req 1.7: Incluir enlace a Grafana con contexto temporal de la anomalía.", GREEN),
    ("Req 2.1: Entregar documentación técnica del modelo, entrenamiento y configuración.", GREEN),
    ("Req 2.2: Entregar memoria del trabajo final con arquitectura y resultados.", YELLOW),
    ("Req 2.3: Incluir guía básica de despliegue e integración.", GREEN),
    ("Req 3.1: Validar el sistema sobre datos históricos representativos del régimen normal.", GREEN),
    ("Req 3.2: Generar métricas de desempeño (falsos positivos, precisión de detección).", GREEN),
    ("Req 3.3: Calibrar el umbral de error con el equipo de operaciones.", GREEN),
    ("Req 4.1: Integrarse con dashboards existentes en Grafana sin GUI propia.", GREEN),
    ("Req 4.2: Incluir en la alerta timestamp, métricas anómalas y valores real vs reconstruido.", GREEN),
    ("Req 5.1: Integrarse con la instancia Prometheus de la empresa.", GREEN),
    ("Req 5.2: Ser compatible con Grafana para visualización.", GREEN),
    ("Req 5.3: Emitir alertas en formato Opsgenie.", GREEN),
    ("Req 6.1: No enviar datos a servicios externos no autorizados.", GREEN),
    ("Req 6.2: Procesar dentro de la red interna o infraestructura autorizada en AWS.", GREEN),
    ("Req 6.3: Respetar políticas internas de privacidad y seguridad.", GREEN),
    ("Req 7.1 (opc.): Incluir is_weekend, is_night y variables temporales adicionales.", GREEN),
    ("Req 7.2 (opc.): Permitir ajuste del umbral desde configuración externa.", GREEN),
]

RISKS = [
    (
        "Riesgo 1: Retrasos por disponibilidad limitada de recursos computacionales en AWS. "
        "S=8, O=6, RPN=48 → S*=6, O*=4, RPN*=24.",
        GREEN,
    ),
    (
        "Riesgo 2: Baja calidad o inconsistencias en los datos históricos de Prometheus. "
        "S=7, O=5, RPN=35 → S*=6, O*=4, RPN*=24.",
        GREEN,
    ),
    (
        "Riesgo 3: Alta tasa de falsos positivos en producción. "
        "S=9, O=4, RPN=36 → S*=6, O*=3, RPN*=18.",
        GREEN,
    ),
    (
        "Riesgo 4: Retrasos en la integración con Grafana u Opsgenie. "
        "S=6, O=5, RPN=30 → S*=5, O*=3, RPN*=15.",
        GREEN,
    ),
    (
        "Riesgo 5: Reentrenamiento frecuente por cambios en el comportamiento del servicio. "
        "S=5, O=6, RPN=30 → S*=4, O*=4, RPN*=16.",
        YELLOW,
    ),
]


def set_cell_shading(cell, color: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(child)
    if color:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)


def set_cell_border(cell, thick: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    sz = "18" if thick else "4"
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def insert_paragraph_after(paragraph, text: str = "", bold: bool = False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def insert_table_after(paragraph, rows: int, cols: int):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    return table


def style_body_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_body_text(after_paragraph, text: str):
    current = after_paragraph
    for block in text.strip().split("\n\n"):
        current = insert_paragraph_after(current, block.strip())
        style_body_paragraph(current)
    return current


def add_heading(after_paragraph, text: str):
    p = insert_paragraph_after(after_paragraph, text, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_image(after_paragraph, image_path: Path, caption: str, width=Inches(5.8)):
    p = insert_paragraph_after(after_paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=width)
    cap = insert_paragraph_after(p, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.italic = True
    return cap


def remove_table(table) -> None:
    table._element.getparent().remove(table._element)


def rebuild_colored_list_table_after(paragraph, items: list[tuple[str, str]]):
    table = insert_table_after(paragraph, rows=len(items), cols=1)
    for idx, (text, color) in enumerate(items):
        cell = table.rows[idx].cells[0]
        cell.text = text
        set_cell_shading(cell, color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    return table


def find_paragraph(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(f"No se encontró párrafo con prefijo: {prefix}")


SECTION1 = """
El presente informe resume el estado del Trabajo Final de Especialización titulado «Detección de anomalías en servicio de TV-over-IP mediante autoencoder LSTM», desarrollado por el Ing. Christopher Charaf bajo la dirección de la Esp. Ing. María Fabiana Cid (FIUBA), en el marco del proyecto profesional en Kaltura Inc. El trabajo propone un sistema de detección de anomalías multivariadas integrado al stack de observabilidad ya adoptado por la organización (Prometheus, Grafana y Opsgenie), manteniendo la telemetría y el entrenamiento del modelo dentro del perímetro controlado por la empresa.

Consistencia del trabajo realizado

A la fecha de este informe (mayo de 2026) se completó la implementación funcional del sistema y la validación operativa con el equipo de monitoreo. El núcleo analítico —pipeline de datos, autoencoder LSTM, calibración de umbral e integración con alertas— se encuentra desplegado y operativo. La memoria técnica documenta en detalle los capítulos de introducción, recursos, diseño e implementación; restan en redacción final los capítulos de ensayos y conclusiones previstos para la entrega del Taller de Trabajo Final en junio de 2026.

El sistema procesa ventanas deslizantes de 20 pasos temporales (10 minutos) con muestreo cada 30 segundos sobre once variables: cinco métricas operativas del servicio (tasa de peticiones, latencia percentil 95, memoria, tasa de errores y CPU) más seis variables de contexto temporal con codificación cíclica e indicadores de fin de semana y horario nocturno. Las consultas PromQL, las cotas de normalización fixed_minmax y los parámetros del modelo se externalizaron en archivos YAML bajo config/, lo que permitió auditar y reproducir el comportamiento sin modificar código fuente.

Entrenamiento y validación

El entrenamiento se realizó sobre 90 días de telemetría histórica obtenida desde la instancia interna de Prometheus de la empresa, reservando el 20 % final de las ventanas para validación temporal cronológica. El autoencoder LSTM adoptó una arquitectura codificador-decodificador con capas de 64, 32 y 16 unidades, función de pérdida MSE y early stopping sobre la pérdida de validación. El umbral de detección se calibró como percentil 99,5 de los errores de reconstrucción en validación y se ajustó iterativamente con operaciones para equilibrar sensibilidad y falsos positivos.

Durante el desarrollo se identificaron y resolvieron dos decisiones críticas de diseño. Primero, se reemplazó StandardScaler por MinMaxScaler con cotas fijas, porque el escalado adaptativo entrenado con datos sintéticos producía falsos positivos masivos al evaluar sobre Prometheus productivo. Segundo, se implementó un servicio simulado y un generador sintético con paridad semántica PromQL para el repositorio académico reproducible, sin sustituir el entrenamiento real autorizado dentro de la infraestructura de la empresa.

Integración operativa

El servicio de inferencia (scripts/inference.py) ejecuta ciclos periódicos de consulta a Prometheus, preprocesamiento, reconstrucción y comparación con el umbral persistido. Ante una anomalía confirmada, el flujo emite alertas HTTPS hacia Opsgenie con prioridad derivada del exceso sobre el umbral, deduplicación temporal, filtrado por confianza mínima y enlace contextual a un panel de Grafana. El despliegue se containerizó con Docker Compose para el perfil de demostración y se integró a la red interna existente en producción, con límites de 1 vCPU y 2 GB de RAM suficientes para el ciclo de 30 segundos.

La validación funcional con el equipo de operaciones confirmó la detección de degradaciones acordadas, una tasa de falsos positivos compatible con la capacidad de respuesta del servicio de guardia y la recepción correcta de alertas y enlaces en Opsgenie y Grafana. El código fuente, los artefactos del modelo (pesos, preprocesador y umbral), la configuración de integración y la documentación de despliegue se encuentran en el repositorio entregable del proyecto.

Perspectiva de cierre

Se estima que el proyecto podrá completarse en su totalidad antes del inicio del Taller de Trabajo Final de junio de 2026. Las actividades pendientes se concentran en la redacción final de los capítulos de ensayos y conclusiones de la memoria, la consolidación de este informe de avances y la preparación de la presentación oral. No se identifican bloqueos técnicos abiertos: el sistema cumple los requerimientos funcionales, de integración y de seguridad definidos en la planificación inicial, y los riesgos de mayor impacto (calidad de datos, falsos positivos e integración externa) fueron mitigados durante la validación con operaciones.
"""


def main() -> None:
    doc = Document(str(DOC_PATH))

    doc.paragraphs[0].text = (
        "Detección de anomalías en servicio de TV-over-IP mediante autoencoder LSTM"
    )
    doc.paragraphs[3].text = "Ing. Christopher Charaf"
    doc.paragraphs[6].text = "Esp. Ing. María Fabiana Cid (FIUBA)"
    doc.paragraphs[17].text = (
        "Este plan de trabajo ha sido realizado en el marco de la asignatura "
        "Gestión de Proyectos entre el 24 de junio y el 20 de octubre de 2025."
    )

    rev_table = doc.tables[0]
    rev_table.rows[1].cells[0].text = "1.0"
    rev_table.rows[1].cells[1].text = "Creación del documento"
    rev_table.rows[1].cells[2].text = "01/09/2018"
    new_rev = rev_table.add_row()
    new_rev.cells[0].text = "2.0"
    new_rev.cells[1].text = "Completado con avance del proyecto (mayo 2026)"
    new_rev.cells[2].text = "22/05/2026"

    consigna = doc.paragraphs[31]
    current = add_body_text(consigna, SECTION1)

    current = add_heading(current, "Arquitectura del sistema")
    current = add_body_text(
        current,
        "La figura siguiente resume el flujo operativo desde la recolección en Prometheus "
        "hasta la emisión de alertas y enlaces a Grafana.",
    )
    flow_img = ROOT / "Figures" / "flow_diag.png"
    if flow_img.exists():
        current = add_image(
            current,
            flow_img,
            "Figura 1. Flujo del sistema de detección de anomalías.",
        )

    current = add_heading(current, "Variables del modelo y entorno de demostración")
    current = add_body_text(
        current,
        "La tabla 1 resume las métricas operativas configuradas en config/data.yaml. "
        "Dado que la telemetría productiva es confidencial, el repositorio académico "
        "emplea un entorno simulado con la misma familia de métricas; el entrenamiento "
        "sobre datos reales se realizó dentro de la infraestructura autorizada de la empresa.",
    )

    table_data = [
        ("request_rate", "sum", "Demanda instantánea sobre el servicio"),
        ("latency_p95", "max", "Cola superior de tiempos de respuesta"),
        ("memory_usage", "mean", "Presión de memoria del proceso"),
        ("error_rate", "sum", "Tasa de respuestas erróneas"),
        ("cpu_usage", "sum", "Consumo de CPU del servicio"),
    ]
    t = insert_table_after(current, rows=1, cols=3)
    hdr = t.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Agregación"
    hdr[2].text = "Interpretación"
    for var, agg, desc in table_data:
        row = t.add_row().cells
        row[0].text = var
        row[1].text = agg
        row[2].text = desc
    cap = insert_paragraph_after(current, "Tabla 1. Métricas operativas del servicio.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.italic = True

    dist_img = ROOT / "Figures" / "dataset_distribution.png"
    if dist_img.exists():
        current = add_image(
            cap,
            dist_img,
            "Figura 2. Distribución ilustrativa de métricas en el entorno de demostración.",
            width=Inches(5.2),
        )

    remove_table(doc.tables[1])
    remove_table(doc.tables[1])

    replace_task_table_after_anchor(doc)

    p_req = find_paragraph(doc, "Si considera que es necesario modificar los requerimientos")
    rebuild_colored_list_table_after(p_req, REQUIREMENTS)

    p_risk = find_paragraph(doc, "Si considera que es necesario modificar los riesgos")
    rebuild_colored_list_table_after(p_risk, RISKS)

    doc.save(str(DOC_PATH))
    print(f"Informe actualizado: {DOC_PATH}")


if __name__ == "__main__":
    main()
