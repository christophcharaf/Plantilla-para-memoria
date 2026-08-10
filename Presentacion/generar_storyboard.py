"""Genera storyboard_defensa.docx con imágenes embebidas."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

base = Path(__file__).resolve().parent.parent
figures = base / "Figures"
out = base / "Presentacion" / "storyboard_defensa.docx"

slides = [
    {
        "num": 1,
        "title": "Portada",
        "text": [
            "Detección de anomalías en servicio de TV-over-IP mediante autoencoder LSTM",
            "",
            "Ing. Christopher Charaf",
            "Carrera de Especialización en Inteligencia Artificial — FIUBA",
            "",
            "Director: Esp. Ing. María Fabiana Cid",
            "Defensa — agosto 2026",
        ],
        "images": [],
        "notes": "Presentación del autor; agradecimientos al jurado y a la directora. (~30 s)",
    },
    {
        "num": 2,
        "title": "El recorrido de hoy",
        "text": [
            "1. ¿Por qué este proyecto?",
            "2. La meta",
            "3. Cómo quedó armado",
            "4. El modelo",
            "5. Resultados",
            "6. Demo en vivo",
            "7. Cierre",
        ],
        "images": [],
        "notes": "Seguir el orden de la memoria; la demo va al final. (~30 s)",
    },
    {
        "num": 3,
        "title": "¿Por qué este proyecto?",
        "text": [
            "• Trabajo final — Especialización en IA (FIUBA)",
            "• Aplicación en plataforma de video empresarial",
            "• Servicio TV-over-IP sobre infraestructura en la nube",
            "• Telemetría en Prometheus (muestreo ~30 s)",
        ],
        "images": [],
        "notes": "Responder: contexto del autor y motivación del trabajo. Mencionar colaboración con operaciones para calibrar el umbral. (~1 min)",
    },
    {
        "num": 4,
        "title": "El problema de las 3 a.m.",
        "text": [
            "• Umbrales fijos: útiles para síntomas claros",
            "• Débiles ante patrones acoplados entre métricas",
            "• Muchas alertas sin incidente real → fatiga operativa",
            "• Telemetría debe quedar dentro del perímetro corporativo",
        ],
        "images": [],
        "notes": "El dolor real es la fatiga de alertas, no solo detectar anomalías. (~1,5 min)",
    },
    {
        "num": 5,
        "title": "Cuando algo no cuadra",
        "text": [
            "• Patrón que se aparta del régimen nominal aprendido",
            "• Puede combinar latencia, tráfico y recursos",
            "• Detección temprana → menos impacto en el usuario",
        ],
        "images": ["anomaly_example.png"],
        "notes": "Señalar banda verde = nominal, tramo rojo = desvío sostenido. Fig. 1.1 de la memoria. (~45 s)",
    },
    {
        "num": 6,
        "title": "¿Qué hace la gente hoy?",
        "text": [
            "Umbrales fijos → regla por métrica; poco contexto multivariado",
            "Clásicos (IF, kNN) → vector de features; mucho ajuste manual",
            "Autoencoder LSTM → reconstruye lo nominal; requiere pipeline",
            "Productos comerciales → integrado al vendor; poco control del modelo",
        ],
        "images": [],
        "notes": "Ubicar la propuesta: AE-LSTM + stack Prometheus–Grafana–Opsgenie. (~1,5 min)",
    },
    {
        "num": 7,
        "title": "La meta del trabajo",
        "text": [
            "General: sistema de detección en operación, integrado al monitoreo",
            "",
            "1. Preparar métricas Prometheus (ventanas multivariadas)",
            "2. Entrenar autoencoder LSTM + umbral calibrado",
            "3. Integrar alertas contextualizadas y enlaces a Grafana",
            "4. Evaluar frente a línea base de umbrales fijos",
        ],
        "images": [],
        "notes": "Los 4 objetivos son verificables; al cierre mostrar que se cumplieron. (~1,5 min)",
    },
    {
        "num": 8,
        "title": "Con qué se armó",
        "text": [
            "ML: TensorFlow / Keras, scikit-learn, pandas",
            "Observabilidad: Prometheus, Grafana, Opsgenie / JSM",
            "Despliegue: Docker, Python, YAML declarativo",
            "Entrenamiento e inferencia en CPU",
        ],
        "images": [],
        "notes": "Cap. 2 comprimido; no listar todas las bibliotecas. (~1 min)",
    },
    {
        "num": 9,
        "title": "Los datos (sin spoilear producción)",
        "text": [
            "• 90 días de régimen nominal (entrenamiento)",
            "• Partición cronológica 80 % / 20 %",
            "• Prueba: 24 h con 3 anomalías inyectadas",
            "• Métricas reales: confidenciales",
            "• Números publicados: entorno demo reproducible",
        ],
        "images": ["dataset_distribution.png"],
        "notes": "Anticipar pregunta del jurado sobre confidencialidad. (~1 min)",
    },
    {
        "num": 10,
        "title": "Cómo quedó armado",
        "text": [
            "Prometheus → Detector (contenedor) → Opsgenie",
            "                ↓",
            "            Grafana (contexto temporal)",
        ],
        "images": ["infra_demo.png", "flow_diag.png"],
        "notes": "Un solo contenedor nuevo en producción; el resto ya existía. (~1,5 min)",
    },
    {
        "num": 11,
        "title": "De Prometheus al modelo",
        "text": [
            "1. Consulta PromQL → series alineadas",
            "2. Ingeniería de variables + hora cíclica",
            "3. Normalización fixed_minmax",
            "4. Ventanas deslizantes (tamaño 20, stride 1)",
            "5. Misma transformación en entrenamiento e inferencia",
        ],
        "images": ["sliding_window.png"],
        "notes": "Destacar paridad entrenamiento/inferencia y cotas fijas. (~1,5 min)",
    },
    {
        "num": 12,
        "title": "El cerebro: autoencoder LSTM",
        "text": [
            "• Entrenamiento no supervisado (solo régimen nominal)",
            "• Encoder → bottleneck → Decoder",
            "• Señal de anomalía: error de reconstrucción (MSE)",
            "• Umbral: percentil 99,5 del error en validación",
        ],
        "images": ["lstm_autoencoder_arch.png"],
        "notes": "Idea Malhotra et al.: si no reconstruye bien, algo raro pasó. (~1,5 min)",
    },
    {
        "num": 13,
        "title": "De un número raro a una alerta útil",
        "text": [
            "• Confirmación: 2 ciclos anómalos consecutivos",
            "• Deduplicación: 1 incidente → 1 apertura + 1 cierre",
            "• Escalamiento si persiste > 30 min",
            "• Enlace Grafana ±15 min alrededor del evento",
        ],
        "images": ["ciclo_vida_alerta.png"],
        "notes": "Aporte directo a la fatiga de alertas (12+ notificaciones → 1). (~1,5 min)",
    },
    {
        "num": 14,
        "title": "Cómo se puso a prueba",
        "text": [
            "Dos planos: reconstrucción (MSE) y clasificación (anomalías inyectadas)",
            "",
            "3 tipos inyectados en 24 h de prueba:",
            "→ Pico de latencia (30 min)",
            "→ Caída de tráfico (15 min)",
            "→ Pico de memoria (60 min)",
        ],
        "images": [],
        "notes": "2861 ventanas de prueba, 9,3 % anómalas. (~1 min)",
    },
    {
        "num": 15,
        "title": "Números que importan",
        "text": [
            "Precisión     0,942",
            "Sensibilidad  0,981",
            "F1-score      0,961",
            "TFP           0,6 %",
        ],
        "images": ["evaluacion_modelo.png"],
        "notes": "Caminar por los 4 paneles: distribución, timeline, matriz, métricas. (~1,5 min)",
    },
    {
        "num": 16,
        "title": "Nominal vs. algo anda mal",
        "text": [
            "Nominal: reconstrucción sigue la señal → error bajo",
            "Anómala: el modelo proyecta régimen habitual → error alto",
        ],
        "images": ["reconstruccion_ejemplo.png"],
        "notes": "Panel derecho: pico de latencia p95 que el modelo no esperaba. (~1 min)",
    },
    {
        "num": 17,
        "title": "Afinando la sensibilidad",
        "text": [
            "• Percentil 99,5: mejor compromiso (F1 = 0,965)",
            "• Percentil 95: más sensibilidad, más falsos positivos",
            "• Percentil 99,9: umbral demasiado alto → pierde anomalías",
        ],
        "images": ["barrido_umbral.png"],
        "notes": "Calibrado con operaciones, no solo por métrica offline. (~1 min)",
    },
    {
        "num": 18,
        "title": "¿Se superan los umbrales fijos?",
        "text": [
            "Umbrales fijos:  Precisión 0,959 | Sensibilidad 0,963 | F1 0,961 | TFP 0,4 %",
            "Autoencoder LSTM: Precisión 0,942 | Sensibilidad 0,981 | F1 0,961 | TFP 0,6 %",
            "",
            "• F1 equivalente en escenario con anomalías marcadas",
            "• AE-LSTM: mayor sensibilidad multivariada",
            "• Valor agregado: contexto + deduplicación operativa",
        ],
        "images": [],
        "notes": "No se gana en todas las métricas; el valor está en el enfoque integrado. (~1 min)",
    },
    {
        "num": 19,
        "title": "Nadie es perfecto (ni el modelo)",
        "text": [
            "• 5 falsos negativos: todos en caída de tráfico",
            "• Se parece a horarios de baja demanda legítima",
            "• Evaluación con anomalías sintéticas inyectadas",
            "• Validación productiva: cualitativa (datos confidenciales)",
        ],
        "images": [],
        "notes": "Mostrar las limitaciones con claridad; conecta con trabajo futuro. (~1 min)",
    },
    {
        "num": 20,
        "title": "A verlo en acción",
        "text": [
            "Flujo extremo a extremo:",
            "1. Anomalía inyectada en servicio simulado",
            "2. Detector marca error elevado",
            "3. Alerta en JSM Operations",
            "4. Operador abre Grafana con contexto",
        ],
        "images": ["demo_grafana.png", "demo_jsm_alerta.png"],
        "notes": "Puente hacia demo en vivo o video (3–5 min). Ver anexo de escenas al final del documento.",
    },
    {
        "num": 21,
        "title": "Qué aporta esto de verdad",
        "text": [
            "• Pipeline reproducible YAML + paridad train/infer",
            "• Escalado fixed_minmax sin recalibrar en cada ciclo",
            "• Señal multivariada única → umbral más simple",
            "• Deduplicación: de 12+ alertas a 1 por incidente",
            "• Integración al stack de observabilidad existente",
        ],
        "images": [],
        "notes": "Respuesta directa a: ¿cuáles son los aportes? (~1,5 min)",
    },
    {
        "num": 22,
        "title": "¿Y ahora qué?",
        "text": [
            "1. Mejorar detección de caídas de tráfico sutiles",
            "2. Retroalimentación con operaciones (FP/FN reales)",
            "3. Reentrenamiento y promoción automatizada de artefactos",
            "4. Explorar arquitecturas alternativas si crece la complejidad",
            "5. Extender a otros servicios del ecosistema TV-over-IP",
        ],
        "images": [],
        "notes": "Tres bullets alcanzan si el tiempo es justo; tener el resto en backup. (~1 min)",
    },
    {
        "num": 23,
        "title": "Lo que aportó la carrera",
        "text": [
            "• ML: RNN/LSTM, aprendizaje no supervisado, métricas",
            "• Datos: series temporales, partición cronológica",
            "• Software: Python, Docker, APIs, configuración declarativa",
            "• Operaciones: observabilidad, gestión de alertas, SRE",
        ],
        "images": [],
        "notes": "Slide rápida; el jurado de IA valora el puente con ops. (~45 s)",
    },
    {
        "num": 24,
        "title": "En resumen...",
        "text": [
            "• Objetivo general cumplido: sistema integrado en operación",
            "• Detección multivariada con baja tasa de falsas alertas",
            "• Demo reproducible y validación cualitativa en producción",
            "",
            "¡Gracias! — Preguntas",
        ],
        "images": [],
        "notes": "Cerrar: el aporte central es llevar ML al monitoreo sin externalizar telemetría. (~30 s)",
    },
]


def add_images(doc, image_names):
    if not image_names:
        return
    doc.add_paragraph().add_run("Imágenes:").bold = True
    for img_name in image_names:
        img_path = figures / img_name
        p = doc.add_paragraph()
        p.add_run(f"Archivo: Figures/{img_name}").italic = True
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(5.8))
                cap = doc.add_paragraph(f"[Vista previa embebida: {img_name}]")
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            except Exception as exc:
                doc.add_paragraph(f"(No se pudo embeber: {exc})")
        else:
            doc.add_paragraph(f"(Archivo no encontrado: {img_path})")


def main():
    doc = Document()

    title = doc.add_heading("Storyboard — Presentación de defensa", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Detección de anomalías en servicio de TV-over-IP mediante autoencoder LSTM"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        "Ing. Christopher Charaf | FIUBA — Especialización en IA | 24 slides + demo"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run("Uso: ").bold = True
    intro.add_run(
        "Guía para armar la presentación en Canva. Cada sección = 1 slide. "
        "Las imágenes están embebidas desde Figures/ de la memoria."
    )

    doc.add_page_break()

    for slide in slides:
        doc.add_heading(f"SLIDE {slide['num']} — {slide['title']}", level=1)

        doc.add_paragraph().add_run("Texto en pantalla:").bold = True
        for line in slide["text"]:
            p = doc.add_paragraph(line if line else " ")
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(11)

        add_images(doc, slide["images"])

        doc.add_paragraph().add_run("Notas de orador:").bold = True
        note_p = doc.add_paragraph(slide["notes"])
        note_p.runs[0].italic = True
        note_p.runs[0].font.size = Pt(10)

        doc.add_page_break()

    doc.add_heading("Anexo — Escenas del video demo (3–5 min)", level=1)
    demo_scenes = [
        ("A", "30 s", "demo_grafana.png", "Pico sostenido de latencia p95"),
        ("B", "45 s", "demo_logs.png", "Detector: error > umbral, confirmación 2 ciclos"),
        ("C", "45 s", "demo_jsm_alerta.png", "Apertura de alerta con enlace a Grafana"),
        ("D", "45 s", "demo_jsm_escalacion.png", "Escalamiento por persistencia"),
        ("E", "30 s", "ciclo_vida_alerta.png", "Resolución y deduplicación"),
    ]
    for scene, dur, img, text in demo_scenes:
        doc.add_heading(f"Escena {scene} ({dur})", level=2)
        doc.add_paragraph(f"Texto overlay: {text}")
        img_path = figures / img
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(5.8))
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Slides backup (solo si preguntan)", level=1)
    backups = [
        ("B1 — Hiperparámetros", "Tabla 3.9: capas LSTM, batch, epochs, etc."),
        ("B2 — Módulos del código", "Tabla 3.4 resumida"),
        ("B3 — Costo de inferencia", "Tabla 3.7: CPU, RAM, ciclo 30 s"),
        ("B4 — Config YAML", "data.yaml, alerting.yaml, etc."),
    ]
    for title_text, detail in backups:
        doc.add_paragraph(title_text, style="List Bullet")
        doc.add_paragraph(detail)

    doc.save(str(out))
    print(f"Created: {out}")
    print(f"Size: {out.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
